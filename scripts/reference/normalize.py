#!/usr/bin/env python3
"""
Нормализация pandoc-reference.docx:
- именованные стили получают шрифт/размер вместо прямого форматирования;
- прямое форматирование рунов очищается;
- создаётся/обновляется стиль "Table" под нужный вид;
- колонтитул заменяется на поле STYLEREF "Heading 1".

Запуск: ./normalize.py [input.docx] [output.docx]
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Pt

BODY_FONT = "Golos Text"
HEADING_FONT = "Raleway Thin"

STYLE_FONTS = {
    "normal": (BODY_FONT, Pt(10)),
    "Heading 1": (HEADING_FONT, Pt(15)),
    "Heading 2": (HEADING_FONT, Pt(12)),
    "Heading 3": (HEADING_FONT, Pt(11)),
    "Heading 4": (HEADING_FONT, Pt(10)),
    "Heading 5": (HEADING_FONT, Pt(10)),
    "Heading 6": (HEADING_FONT, Pt(10)),
    "Title": (HEADING_FONT, Pt(22)),
    "Subtitle": (HEADING_FONT, Pt(14)),
}


def set_style_font(style, name, size):
    f = style.font
    f.name = name
    f.size = size
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rfonts.set(qn(f"w:{attr}"), name)


def strip_direct_run_formatting(paragraph):
    for run in paragraph.runs:
        run.font.name = None
        run.font.size = None
        rpr = run._r.find(qn("w:rPr"))
        if rpr is not None:
            for tag in ("w:rFonts", "w:sz", "w:szCs"):
                el = rpr.find(qn(tag))
                if el is not None:
                    rpr.remove(el)


def ensure_compact_style(doc):
    """Pandoc использует стиль 'Compact' для абзацев в ячейках таблиц.
    Если его нет — содержимое ячеек рендерится без форматирования."""
    name = "Compact"
    try:
        st = doc.styles[name]
    except KeyError:
        st = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        st.base_style = doc.styles["normal"]
    f = st.font
    f.name = BODY_FONT
    f.size = Pt(10)
    pf = st.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    return st


HEADER_BG = "434343"
HEADER_FG = "ffffff"
HEADER_BORDER = "1a56a0"
CELL_BORDER = "bbcfe8"


def _make_borders(color):
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), color)
        borders.append(b)
    return borders


def _replace_child(parent, tag, new_el):
    existing = parent.find(qn(tag))
    if existing is not None:
        parent.remove(existing)
    parent.append(new_el)


def ensure_table_style(doc):
    name = "Table"
    styles = doc.styles
    try:
        st = styles[name]
    except KeyError:
        st = styles.add_style(name, WD_STYLE_TYPE.TABLE)

    f = st.font
    f.name = BODY_FONT
    f.size = Pt(10)

    el = st.element
    tbl_pr = el.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        el.append(tbl_pr)
    _replace_child(tbl_pr, "w:tblBorders", _make_borders(CELL_BORDER))

    # Conditional formatting for the header row — pandoc sets tblLook firstRow=1.
    for old in el.findall(qn("w:tblStylePr")):
        el.remove(old)
    hdr = OxmlElement("w:tblStylePr")
    hdr.set(qn("w:type"), "firstRow")

    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rfonts.set(qn(f"w:{attr}"), BODY_FONT)
    rpr.append(rfonts)
    bold = OxmlElement("w:b")
    bold.set(qn("w:val"), "1")
    rpr.append(bold)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), HEADER_FG)
    rpr.append(color)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "20")
    rpr.append(sz)
    hdr.append(rpr)

    ppr = OxmlElement("w:pPr")
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    ppr.append(jc)
    hdr.append(ppr)

    h_tcpr = OxmlElement("w:tcPr")
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), HEADER_BG)
    h_tcpr.append(shd)
    valign = OxmlElement("w:vAlign")
    valign.set(qn("w:val"), "center")
    h_tcpr.append(valign)
    h_borders = _make_borders(HEADER_BORDER)
    h_tcpr.append(h_borders)
    hdr.append(h_tcpr)

    el.append(hdr)
    return st


def apply_table_style(doc, style_name="Table"):
    for t in doc.tables:
        t.style = doc.styles[style_name]


def replace_header_with_styleref(doc, style_name="Heading 1"):
    for sec in doc.sections:
        hdr = sec.header
        for p in list(hdr.paragraphs):
            for r in list(p.runs):
                r._r.getparent().remove(r._r)
            run = p.add_run()
            fld_begin = OxmlElement("w:fldChar")
            fld_begin.set(qn("w:fldCharType"), "begin")
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = f' STYLEREF "{style_name}" \\* MERGEFORMAT '
            fld_sep = OxmlElement("w:fldChar")
            fld_sep.set(qn("w:fldCharType"), "separate")
            cached = OxmlElement("w:t")
            cached.text = ""
            fld_end = OxmlElement("w:fldChar")
            fld_end.set(qn("w:fldCharType"), "end")
            for el in (fld_begin, instr, fld_sep, cached, fld_end):
                run._r.append(el)
            return


def replace_body_with_stub(doc):
    """Pandoc читает из reference-doc только определения стилей,
    содержимое body не используется. Заменяем на нейтральную рыбу."""
    body = doc.element.body
    sectPr = body.find(qn("w:sectPr"))
    for child in list(body):
        if child is not sectPr:
            body.remove(child)

    def _add(text, style):
        p = doc.add_paragraph(text, style=doc.styles[style])
        return p

    _add("Заголовок документа", "Title")
    _add("Раздел первого уровня", "Heading 1")
    _add("Подраздел", "Heading 2")
    _add("Подподраздел", "Heading 3")
    _add(
        "Основной текст документа набирается стилем «normal». "
        "Для проверки начертания: жирный, курсив, обычный.",
        "normal",
    )

    table = doc.add_table(rows=2, cols=3)
    table.style = doc.styles["Table"]
    hdr = table.rows[0].cells
    hdr[0].text = "Колонка A"
    hdr[1].text = "Колонка B"
    hdr[2].text = "Колонка C"
    row = table.rows[1].cells
    row[0].text = "Значение"
    row[1].text = "Значение"
    row[2].text = "Значение"

    if sectPr is not None:
        body.remove(sectPr)
        body.append(sectPr)


def main():
    here = Path(__file__).parent
    default = here / "pandoc-reference.docx"
    inp = Path(sys.argv[1]) if len(sys.argv) > 1 else default
    out = Path(sys.argv[2]) if len(sys.argv) > 2 else default

    doc = Document(str(inp))

    for style_name, (font, size) in STYLE_FONTS.items():
        try:
            set_style_font(doc.styles[style_name], font, size)
        except KeyError:
            print(f"  [skip] стиль не найден: {style_name}")

    for p in doc.paragraphs:
        strip_direct_run_formatting(p)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    strip_direct_run_formatting(p)

    ensure_compact_style(doc)
    ensure_table_style(doc)
    apply_table_style(doc)
    replace_header_with_styleref(doc, "Heading 1")
    replace_body_with_stub(doc)

    doc.save(str(out))
    print(f"Готово: {out}")


if __name__ == "__main__":
    main()
